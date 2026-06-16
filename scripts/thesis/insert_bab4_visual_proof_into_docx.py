"""Insert maximum-detail visual proof evidence into BAB IV DOCX.

This pass intentionally avoids Word tables in the output. Existing Word tables
are rendered into PNG images and the tables are removed from the generated copy.
The source DOCX is not modified.
"""

from __future__ import annotations

import argparse
import json
import textwrap
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt


DEFAULT_SOURCE = Path(
    r"C:\Users\ACER\Downloads\Penelitian\TA\IMPLEMENTASI HYBRID MODEL COLLABORATIVE FILTERING - BAB IV SCPA Full Final + Evidence + Benchmark + DQN Proxy.docx"
)
DEFAULT_OUTPUT = Path(
    r"C:\Users\ACER\Downloads\Penelitian\TA\IMPLEMENTASI HYBRID MODEL COLLABORATIVE FILTERING - BAB IV SCPA Full Final + Visual Proof Max.docx"
)
REPO_ROOT = Path(__file__).resolve().parents[2]
LATEST_VISUAL = REPO_ROOT / "reports" / "thesis_evidence" / "bab4_visual_proof_latest.json"


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def usable_width_inches(doc: Document) -> float:
    section = doc.sections[0]
    width_emu = section.page_width - section.left_margin - section.right_margin
    return max(3.5, min(6.1, width_emu / 914400))


def insert_paragraph(doc: Document, ref_para, text: str = "", style: str | None = None):
    para = doc.add_paragraph(text, style=style)
    ref_para._p.addprevious(para._p)
    return para


def insert_image(doc: Document, ref_para, image_path: Path, caption: str, width_inches: float) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    ref_para._p.addprevious(para._p)
    cap = insert_paragraph(doc, ref_para, caption, "Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)


def insert_body(doc: Document, ref_para, text: str) -> None:
    para = insert_paragraph(doc, ref_para, text)
    para.paragraph_format.space_after = Pt(6)


def wrap_cell(value: str, width: int) -> str:
    value = " ".join(str(value or "").split())
    return "\n".join(textwrap.wrap(value, width=width, break_long_words=False)) or " "


def table_to_rows(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    # Drop fully empty trailing rows.
    return [row for row in rows if any(cell.strip() for cell in row)]


def render_word_table_image(rows: list[list[str]], path: Path, index: int) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    col_count = max((len(row) for row in rows), default=1)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    header = normalized[0] if normalized else [""]
    body = normalized[1:] if len(normalized) > 1 else []

    # Make long-text tables readable on A5 by wrapping aggressively and lowering
    # font size only when row count demands it.
    width_by_cols = {1: 62, 2: 42, 3: 30, 4: 24, 5: 20, 6: 18}
    wrap_width = width_by_cols.get(col_count, 16)
    display = [[wrap_cell(cell, wrap_width) for cell in row] for row in normalized]
    row_count = max(1, len(display))
    font_size = 8 if row_count <= 6 else 7 if row_count <= 10 else 6
    fig_height = min(6.4, max(1.7, 0.42 * row_count + 0.65))
    fig_width = 8.2

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    fig.patch.set_facecolor("white")
    ax.axis("off")
    ax.text(
        0.0,
        1.02,
        f"Visualisasi tabel Word #{index:02d}",
        transform=ax.transAxes,
        fontsize=9,
        color="#334155",
        fontweight="bold",
        va="bottom",
    )
    table_obj = ax.table(
        cellText=display[1:] if body else display,
        colLabels=header if body else None,
        cellLoc="left",
        colLoc="left",
        loc="upper left",
        bbox=[0, 0, 1, 0.92],
    )
    table_obj.auto_set_font_size(False)
    table_obj.set_fontsize(font_size)
    table_obj.scale(1.0, 1.35)
    for (row, _col), cell in table_obj.get_celld().items():
        cell.set_edgecolor("#1f2937")
        cell.set_linewidth(0.55)
        if body and row == 0:
            cell.set_facecolor("#e8eef7")
            cell.set_text_props(weight="bold", color="#111827")
        else:
            cell.set_facecolor("#ffffff" if row % 2 else "#f8fafc")
            cell.set_text_props(color="#111827")
    fig.tight_layout()
    fig.savefig(path, dpi=260, bbox_inches="tight")
    plt.close(fig)


def convert_word_tables_to_images(doc: Document, table_dir: Path, width_inches: float) -> int:
    count = 0
    # Snapshot the list because we remove tables as we go.
    for table in list(doc.tables):
        rows = table_to_rows(table)
        if not rows:
            parent = table._tbl.getparent()
            parent.remove(table._tbl)
            continue
        count += 1
        image_path = table_dir / f"word_table_{count:03d}.png"
        render_word_table_image(rows, image_path, count)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        table._tbl.addprevious(para._p)
        parent = table._tbl.getparent()
        parent.remove(table._tbl)
    return count


def main() -> int:
    parser = argparse.ArgumentParser(description="Insert visual proof evidence into BAB IV DOCX")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--manifest", type=Path, default=LATEST_VISUAL)
    args = parser.parse_args()

    manifest = read_json(args.manifest)
    assets = {key: Path(value) for key, value in manifest["assets"].items()}
    summary = manifest["summary"]

    doc = Document(str(args.source))
    refs = [p for p in doc.paragraphs if p.text.strip().upper() == "DAFTAR PUSTAKA BAB IV"]
    if not refs:
        raise RuntimeError("Target paragraph 'DAFTAR PUSTAKA BAB IV' not found")
    ref = refs[0]
    width = usable_width_inches(doc)

    page_break = insert_paragraph(doc, ref)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    insert_paragraph(doc, ref, "4.19 Validasi Visual Evidence Maksimal BAB IV", "Heading 2")
    insert_body(
        doc,
        ref,
        (
            "Bagian ini ditambahkan untuk memperkuat BAB IV agar tidak berhenti pada uraian teoritis. "
            "Semua bukti disajikan sebagai gambar: plot matplotlib, screenshot runtime yang sudah "
            "terdokumentasi, terminal bergaya PowerShell, visual card teknis model, dan sumber Mermaid. "
            "Data yang dipakai berasal dari artifact runtime 16 Juni 2026, benchmark frozen split, source "
            "code service, dan qrels status. Jika suatu bukti belum tersedia, statusnya ditulis sebagai "
            "blocker agar klaim akademik tetap sesuai fakta."
        ),
    )
    insert_image(doc, ref, assets["evidence_dashboard"], "Gambar 4.32 Dashboard visual evidence runtime dan benchmark SCPA.", width)
    insert_body(
        doc,
        ref,
        (
            f"Hasil dashboard menunjukkan artifact yang terukur: {summary['runtime_active_accepted_jobs']} lowongan aktif-terverifikasi, "
            f"{summary['runtime_ready_embeddings']} embedding siap dengan dimensi {summary['embedding_dimension']}, "
            f"{summary['benchmark_users']} user benchmark, {summary['benchmark_jobs']} lowongan benchmark, dan "
            f"{summary['benchmark_interactions']} interaksi offline. Nilai full_scpa mencapai NDCG@10 "
            f"{summary['temporal_full_ndcg10']:.4f} pada temporal split dan {summary['holdout_full_ndcg10']:.4f} pada user holdout."
        ),
    )
    insert_body(
        doc,
        ref,
        (
            "Batas interpretasinya adalah angka benchmark berasal dari simulated_grounded, yaitu simulasi "
            "perilaku yang diturunkan dari atribut nyata profil dan lowongan. Karena itu angka tersebut "
            "valid sebagai evaluasi offline awal, tetapi tidak boleh ditulis sebagai bukti perilaku pengguna produksi."
        ),
    )

    insert_paragraph(doc, ref, "4.19.1 Bukti Visual SBERT Candidate Generator", "Heading 3")
    insert_image(doc, ref, assets["sbert_card"], "Gambar 4.33 Rincian teknis visual SBERT candidate generator.", width)
    insert_body(
        doc,
        ref,
        (
            "SBERT pada SCPA dibuktikan sebagai candidate generator, bukan model yang langsung mengambil "
            "keputusan akhir. Input teks lowongan dibentuk oleh canonical_job_text, embedding disimpan sebagai "
            "vector 384 dimensi, dan retrieval memakai cosine similarity pada pgvector. Contoh response API "
            "membuktikan sbert_score ikut keluar dalam kontrak rekomendasi."
        ),
    )
    insert_image(doc, ref, assets["database_embedding_coverage"], "Gambar 4.34 Plot coverage database embedding lowongan.", width)
    insert_image(doc, ref, assets["qrels_status"], "Gambar 4.35 Status visual qrels SBERT dan blocker gold annotation.", width)
    insert_body(
        doc,
        ref,
        (
            "Batas interpretasinya adalah SBERT boleh diklaim sudah terintegrasi secara runtime dan memiliki "
            "silver qrels, tetapi belum boleh diklaim memiliki Precision, MAP, atau NDCG berbasis expert sebelum "
            "gold qrels dan inter-annotator agreement selesai."
        ),
    )

    insert_paragraph(doc, ref, "4.19.2 Bukti Visual NCF Personalization Scorer", "Heading 3")
    insert_image(doc, ref, assets["ncf_card"], "Gambar 4.36 Rincian teknis visual NCF personalization scorer.", width)
    insert_image(doc, ref, assets["split_counts"], "Gambar 4.37 Plot split train, validation, dan test untuk NCF/hybrid.", width)
    insert_body(
        doc,
        ref,
        (
            "NCF dibuktikan sebagai personalization scorer karena menerima representasi user-item dan belajar "
            "dari bobot event implicit feedback. Pengujian offline memakai temporal split dan user holdout agar "
            "risiko leakage dapat dibaca, bukan hanya diasumsikan. Hasil NCF-only juga sengaja ditampilkan dalam "
            "ablation untuk menunjukkan bahwa sinyal collaborative tanpa kandidat semantik belum cukup kuat."
        ),
    )
    insert_body(
        doc,
        ref,
        (
            "Batas interpretasinya adalah personalization gain pada BAB IV ini tetap terikat pada benchmark "
            "offline. Data real feedback runtime masih berada pada level readiness smoke sehingga belum cukup "
            "untuk klaim generalisasi perilaku pengguna."
        ),
    )

    insert_paragraph(doc, ref, "4.19.3 Bukti Visual DQN Session Reranker", "Heading 3")
    insert_image(doc, ref, assets["dqn_card"], "Gambar 4.38 Rincian teknis visual DQN session reranker.", width)
    insert_image(doc, ref, assets["dqn_session_delta_ndcg"], "Gambar 4.39 Histogram delta NDCG@10 DQN pada held-out session.", width)
    insert_image(doc, ref, assets["dqn_rank_delta_examples"], "Gambar 4.40 Contoh rank_before_dqn dan rank_after_dqn dari event sesi.", width)
    insert_image(doc, ref, assets["dqn_rank_before_after_scatter"], "Gambar 4.41 Scatter rank_before_dqn versus rank_after_dqn.", width)
    insert_body(
        doc,
        ref,
        (
            f"DQN aktif setelah kandidat SBERT dan skor NCF tersedia. Evidence runtime dan benchmark menampilkan "
            f"rank_before_dqn, rank_after_dqn, dqn_session_score, dan event sesi. Stabilitas policy memiliki CV "
            f"{summary['dqn_stability_cv']:.4f}, sedangkan proxy held-out session menghasilkan delta NDCG@10 "
            f"{summary['dqn_proxy_delta_ndcg10']:.4f}."
        ),
    )
    insert_body(
        doc,
        ref,
        (
            "Batas interpretasinya adalah DQN tidak dipakai sebagai learning path. State, action, dan reward "
            "diarahkan untuk reranking slate lowongan dalam satu sesi, bukan untuk menyusun jalur belajar, roadmap "
            "kompetensi, atau urutan karier adaptif."
        ),
    )

    insert_paragraph(doc, ref, "4.19.4 Bukti Visual Hybrid dan Ablation", "Heading 3")
    insert_image(doc, ref, assets["hybrid_card"], "Gambar 4.42 Evidence hybrid, ablation, dan batas klaim.", width)
    insert_image(doc, ref, assets["ablation_ndcg10"], "Gambar 4.43 Plot ablation NDCG@10 SBERT, NCF, DQN, dan full_scpa.", width)
    insert_image(doc, ref, assets["dqn_stability"], "Gambar 4.44 Plot stabilitas DQN multi-seed.", width)
    insert_body(
        doc,
        ref,
        (
            "Ablation memisahkan kontribusi popularity, content, NCF, content+NCF, dan full_scpa pada data yang sama. "
            "Dengan format ini, pembaca dapat melihat bahwa SCPA bukan hanya pipeline yang dijelaskan, tetapi artifact "
            "yang diuji melalui split, metrik, dan perbandingan varian."
        ),
    )
    insert_body(
        doc,
        ref,
        (
            "Batas interpretasinya adalah delta kecil harus dibaca hati-hati. Jika signifikansi tidak kuat pada split "
            "tertentu, BAB IV harus menulisnya sebagai kontribusi terbatas, bukan sebagai klaim superioritas mutlak."
        ),
    )

    insert_paragraph(doc, ref, "4.19.5 Bukti Visual Runtime Frontend, API, dan Terminal", "Heading 3")
    insert_body(
        doc,
        ref,
        (
            "Screenshot frontend berasal dari full-page Playwright capture. Untuk menjaga keterbacaan pada halaman A5, "
            "gambar yang dimasukkan pada subbagian ini memakai crop bagian atas dari screenshot asli. File full-page "
            "tetap disimpan pada folder evidence sehingga dapat diaudit ulang."
        ),
    )
    screenshot_width = width * 0.9
    insert_image(doc, ref, assets["frontend_recommendations_initial_crop"], "Gambar 4.45 Screenshot Playwright rekomendasi awal dari capture runtime 16 Juni 2026.", screenshot_width)
    insert_image(doc, ref, assets["frontend_recommendations_after_events_crop"], "Gambar 4.46 Screenshot Playwright setelah event sesi pengguna.", screenshot_width)
    insert_image(doc, ref, assets["frontend_model_panel_live"], "Gambar 4.47 Screenshot panel skor model pada frontend.", width)
    insert_image(doc, ref, assets["frontend_api_recommendation_summary"], "Gambar 4.48 Plot ringkasan response API rekomendasi.", width)
    insert_image(doc, ref, assets["powershell_frontend_playwright"], "Gambar 4.49 Screenshot PowerShell ringkasan Playwright runtime.", width)
    insert_image(doc, ref, assets["current_capture_terminal"], "Gambar 4.50 Screenshot PowerShell status capture ulang saat revisi.", width)
    insert_body(
        doc,
        ref,
        (
            "Screenshot frontend membuktikan artifact pernah dijalankan melalui browser automation dan response API "
            "menghasilkan rekomendasi dengan lineage model. Pada saat revisi ini frontend dan gateway tidak listening, "
            "sehingga capture baru tidak dipaksakan. Dokumen tetap membedakan bukti runtime yang sudah ada dan status "
            "recapture yang memerlukan stack aktif."
        ),
    )

    insert_paragraph(doc, ref, "4.19.6 Bukti Command, Notebook, dan Diagram Mermaid", "Heading 3")
    insert_image(doc, ref, assets["terminal_thesis_benchmark"], "Gambar 4.51 Screenshot PowerShell command benchmark thesis.", width)
    insert_image(doc, ref, assets["terminal_gold_qrels"], "Gambar 4.52 Screenshot PowerShell command builder gold qrels.", width)
    insert_image(doc, ref, assets["mermaid_runtime_flow_source"], "Gambar 4.53 Screenshot sumber Mermaid flow rekomendasi runtime.", width)
    insert_body(
        doc,
        ref,
        (
            "Notebook 07_bab4_visual_proof_evidence.ipynb menjadi indeks bukti visual yang dapat dibuka ulang. "
            "Sumber Mermaid juga disimpan sebagai file .mmd sehingga diagram dapat dirender ulang dengan mermaid-cli "
            "jika dibutuhkan oleh pembimbing, sementara DOCX tetap memuat screenshot sumber agar evidence terbaca."
        ),
    )

    table_dir = Path(manifest["run_dir"]) / "converted_word_tables"
    converted = convert_word_tables_to_images(doc, table_dir, width)
    doc.save(str(args.output))
    print(
        json.dumps(
            {
                "status": "ok",
                "output": str(args.output),
                "source": str(args.source),
                "manifest": str(args.manifest),
                "converted_word_tables": converted,
                "converted_table_dir": str(table_dir),
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
