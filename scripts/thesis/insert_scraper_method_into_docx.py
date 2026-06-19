"""Insert scraper method/evidence text into the BAB IV DOCX.

This is a surgical DOCX edit: it preserves the existing BAB IV document and
writes a new output file with scraper explanation inserted into sections 4.2.2
and 4.5.1.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SOURCE = ROOT / "docs" / "thesis" / "bab4" / "BAB_IV_HASIL_DAN_PEMBAHASAN.docx"
DEFAULT_OUTPUT = (
    ROOT
    / "docs"
    / "thesis"
    / "bab4"
    / "BAB_IV_HASIL_DAN_PEMBAHASAN_DENGAN_SCRAPING_DATASET_JELAS.docx"
)

INTERACTION_DATASET_PARAGRAPHS = [
    (
        "Dataset utama pada subbab ini adalah simulated_grounded, yaitu paket "
        "interaksi luring pada direktori data/eval/synthetic/. Paket ini "
        "terdiri atas interactions.jsonl, sessions.jsonl, dan "
        "benchmark_metadata.json. Dataset ini digunakan untuk evaluasi NCF, "
        "DQN session reranker, dan ablation hybrid karena log produksi belum "
        "cukup besar. Interaksi di dalamnya bukan klik pengguna nyata; "
        "peristiwa click, view, save, view_10s, skip, dan apply dibangkitkan "
        "oleh model klik terkontrol dengan seed 42."
    ),
    (
        "Grounding dataset berasal dari atribut nyata profil dan lowongan yang "
        "sudah ada di sistem: kecocokan domain, occupation_group atau kelompok "
        "okupasi, dan irisan keterampilan. Metadata menetapkan bobot afinitas "
        "skill 0,45, domain 0,35, dan occupation 0,20; event positif "
        "didefinisikan sebagai apply, save, click, dan view_10s. Paket "
        "simulated_grounded memuat 14.400 peristiwa dari 300 pengguna sintetis "
        "terhadap 3.818 lowongan, 900 sesi, dan positive_rate 0,4367. Dengan "
        "demikian, dataset ini adalah simulasi berbasis grounding: fitur "
        "pembentuk preferensi berasal dari data profil dan lowongan nyata, "
        "tetapi label interaksinya tetap simulatif."
    ),
    (
        "Sebagai pembanding, dataset real_runtime diekspor dari tabel runtime "
        "feedback_events dan disimpan pada direktori data/eval/real_runtime/. "
        "Paket ini terdiri atas interactions.jsonl, sessions.jsonl, "
        "profiles.jsonl, jobs.jsonl, dan metadata.json. Paket ini berisi 693 "
        "peristiwa dari 10 pengguna dan 213 lowongan yang "
        "diekspor; metadata evaluasi mencatat 212 lowongan yang muncul pada "
        "interaksi. Status bukti masih insufficient_for_generalization karena "
        "jumlah pengguna 10 masih di bawah ambang minimum 30 pengguna. Oleh "
        "sebab itu, real_runtime hanya dipakai sebagai bukti pilot/runtime "
        "awal, bukan dasar klaim generalisasi pengguna nyata."
    ),
]


def normalize(text: str) -> str:
    return " ".join((text or "").split())


def find_paragraph_index(doc: Document, startswith: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if normalize(paragraph.text).startswith(startswith):
            return index
    raise ValueError(f"Paragraph starting with {startswith!r} not found")


def insert_after(paragraph, text: str):
    new_paragraph = paragraph.insert_paragraph_before(text)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def apply_body_style(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Cambria"
        run.font.size = Pt(11)


def insert_scraper_text(doc: Document) -> None:
    # Insert implementation detail after the existing 4.2.2 opening paragraph.
    h422_index = find_paragraph_index(doc, "4.2.2 Implementasi Scraper dan Quality Gate")
    anchor_422 = doc.paragraphs[h422_index + 1]
    inserted = [
        (
            "Secara teknis, logika scraping didokumentasikan pada notebook "
            "services/scraper/scraper.ipynb. Notebook tersebut menunjukkan "
            "bahwa parser menggunakan BeautifulSoup untuk membaca elemen "
            "kandidat lowongan dari selector umum seperti [data-job], .job, "
            ".job-card, .job-listing, .vacancy, article, dan li. Dari setiap "
            "kandidat, sistem mengambil atribut title, company, location, "
            "description, tags, source_url, dan content_hash."
        ),
        (
            "Proses pembersihan dilakukan melalui fungsi clean_text() agar "
            "spasi berlebih dan baris kosong tidak terbawa ke data lowongan. "
            "Fungsi first_text() memilih nilai pertama yang tersedia dari "
            "beberapa selector alternatif, sedangkan tag_texts() mengambil "
            "sinyal skill atau tag dari elemen .tag, .tags li, [data-tag], "
            ".chip, dan .badge. Dengan cara ini, scraper tetap dapat membaca "
            "struktur HTML yang tidak selalu identik antar sumber."
        ),
        (
            "Deduplikasi dilakukan dengan content_hash yang dibentuk dari "
            "kombinasi title, company, dan location. Jika hash yang sama sudah "
            "muncul dalam satu proses ekstraksi, entri berikutnya dianggap "
            "duplikat dan tidak diteruskan sebagai kandidat baru. Batas klaim "
            "dari notebook ini adalah pembuktian parser, normalisasi teks, "
            "ekstraksi field, dan deduplikasi; klaim jumlah lowongan real "
            "tetap harus merujuk pada endpoint /scrape/run dan tabel jobs pada "
            "runtime."
        ),
    ]
    current = anchor_422
    for text in inserted:
        current = insert_after(current, text)
        apply_body_style(current)

    # Insert notebook evidence after the existing 4.5.1 opening paragraph.
    h451_index = find_paragraph_index(doc, "4.5.1 Hasil Pengujian Scraper")
    anchor_451 = doc.paragraphs[h451_index + 1]
    evidence = (
        "Validasi notebook menggunakan contoh HTML berisi tiga kartu lowongan: "
        "dua lowongan unik dan satu lowongan duplikat. Hasil eksekusi "
        "menunjukkan count = 2 dan deduplicated = 1. Assertion pada notebook "
        "menghasilkan pesan All scraper notebook assertions passed, sehingga "
        "skenario uji dasar untuk ekstraksi field, cleaning, dan deduplikasi "
        "berjalan sesuai rancangan."
    )
    inserted_evidence = insert_after(anchor_451, evidence)
    apply_body_style(inserted_evidence)


def clarify_interaction_dataset_text(doc: Document) -> None:
    h443_index = find_paragraph_index(doc, "4.4.3 Dataset Interaksi Pengguna")
    first_paragraph = doc.paragraphs[h443_index + 1]
    second_paragraph = doc.paragraphs[h443_index + 2]

    first_paragraph.text = INTERACTION_DATASET_PARAGRAPHS[0]
    apply_body_style(first_paragraph)
    second_paragraph.text = INTERACTION_DATASET_PARAGRAPHS[1]
    apply_body_style(second_paragraph)

    third_paragraph = insert_after(second_paragraph, INTERACTION_DATASET_PARAGRAPHS[2])
    apply_body_style(third_paragraph)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", default=str(DEFAULT_SOURCE))
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT))
    args = parser.parse_args()

    source = Path(args.source)
    output = Path(args.output)
    if not source.exists():
        raise FileNotFoundError(source)

    doc = Document(str(source))
    insert_scraper_text(doc)
    clarify_interaction_dataset_text(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(output)


if __name__ == "__main__":
    main()
